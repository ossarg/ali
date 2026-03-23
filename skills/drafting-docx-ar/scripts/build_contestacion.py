#!/usr/bin/env python3
"""
build_contestacion.py — Genera DOCX de contestación/citación en garantía para Libra.
Formato fiel a los escritos reales: justificado, encabezados BOLD+UNDERLINE, Arial 11pt.
Márgenes: top=5cm, bot=2cm, left=5cm, right=1.5cm (estilo Mariana Díaz / escritos Libra).

Uso: python3 build_contestacion.py <datos.json> <output.docx>
"""

import json, sys, os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Arial"
SIZE = Pt(11)
JUSTIFY = WD_ALIGN_PARAGRAPH.JUSTIFY


def _run(para, text, bold=False, underline=False, italic=False):
    """Agrega un run con formato al párrafo."""
    run = para.add_run(text)
    run.bold = bold
    run.underline = underline
    run.italic = italic
    run.font.name = FONT
    run.font.size = SIZE
    return run


def _para(doc, align=JUSTIFY, space_before=0, space_after=6, left_indent=None, first_line=None):
    """Crea un párrafo vacío con formato base."""
    p = doc.add_paragraph()
    p.alignment = align
    fmt = p.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    if left_indent is not None:
        fmt.left_indent = Cm(left_indent)
    if first_line is not None:
        fmt.first_line_indent = Cm(first_line)
    return p


def add_comment(doc_element, comment_text, author="Ali — Pipeline"):
    """
    Agrega un comentario de Word en el margen del párrafo.
    Los comentarios aparecen visibles en el margen en MS Word.
    """
    try:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import random, datetime

        comment_id = str(random.randint(1, 9999))
        
        # Build comment XML
        comments_part = None
        try:
            comments_part = doc_element.part.comments_part
        except AttributeError:
            pass
        
        if comments_part is None:
            return  # Skip silently if not supported
            
        comment = OxmlElement('w:comment')
        comment.set(qn('w:id'), comment_id)
        comment.set(qn('w:author'), author)
        comment.set(qn('w:date'), datetime.datetime.now().isoformat())
        comment.set(qn('w:initials'), 'AI')
        cp = OxmlElement('w:p')
        cr = OxmlElement('w:r')
        ct = OxmlElement('w:t')
        ct.text = comment_text
        cr.append(ct)
        cp.append(cr)
        comment.append(cp)
        comments_part._element.append(comment)
        
        # Mark in paragraph
        ref = OxmlElement('w:commentReference')
        ref.set(qn('w:id'), comment_id)
        doc_element._p.append(ref)
    except Exception:
        pass  # Never fail silently on comments


def setup_document(doc):
    """Configura márgenes y estilo base."""
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT
    font.size = SIZE
    
    for section in doc.sections:
        section.top_margin = Cm(5.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(5.0)
        section.right_margin = Cm(1.5)


def section_title(doc, text, numero=None, space_before=10):
    """Encabezado de sección: BOLD + UNDERLINE, justificado. Ej: 'PERSONERIA' o 'I. PERSONERÍA'"""
    p = _para(doc, space_before=space_before, space_after=4)
    label = f"{numero}.\t{text}" if numero else text
    _run(p, label, bold=True, underline=True)
    return p


def body(doc, text, space_before=0, space_after=6, indent=None, comment=None):
    """Párrafo de cuerpo, justificado, Arial 11pt."""
    p = _para(doc, space_before=space_before, space_after=space_after, left_indent=indent)
    _run(p, text)
    if comment:
        add_comment(p, comment)
    return p


def to_roman(n):
    vals = [(10,'X'),(9,'IX'),(5,'V'),(4,'IV'),(1,'I')]
    result = ""
    for val, sym in vals:
        while n >= val:
            result += sym
            n -= val
    return result


# ── Secciones ─────────────────────────────────────────────────────────────────

def sec_encabezado(doc, d):
    tipo = d.get("tipo_escrito", "CONTESTA DEMANDA CITADA EN GARANTÍA – OFRECE PRUEBA")
    p = _para(doc, space_after=8)
    _run(p, tipo, bold=True, underline=True)

    p2 = _para(doc, space_after=8)
    _run(p2, "Señor Juez:")

    letrado = d.get("letrado", {})
    nombre = letrado.get("nombre", "[NOMBRE DEL LETRADO — A COMPLETAR]")
    tomo   = letrado.get("tomo", "__")
    folio  = letrado.get("folio", "__")
    cuit   = letrado.get("cuit", "[CUIT]")
    mail   = letrado.get("mail", "[MAIL — A COMPLETAR]")
    cuit_e = letrado.get("cuit_electronico", cuit)
    caratula  = d.get("caratula", "[CARÁTULA — A COMPLETAR]")
    expediente = d.get("expediente", "[N° EXPEDIENTE — A COMPLETAR]")

    p3 = _para(doc, space_after=8)
    _run(p3, f"{nombre} (T{tomo} F{folio} CPACF, CUIT {cuit}, Monotributista), "
             f"letrada apoderada, mail: {mail}, constituyendo domicilio legal en la calle "
             f"Olazabal 1515, piso 5 of. 505, CABA, y domicilio electrónico {cuit_e}, "
             f"en los autos caratulados: ", bold=True)
    _run(p3, f'"{caratula}"', bold=True)
    _run(p3, f" (Expediente N° {expediente}), a V.S. me presento y digo:", bold=True)


def sec_personeria(doc, d, n):
    section_title(doc, "PERSONERÍA", numero=to_roman(n))
    p = _para(doc, space_before=2, space_after=4)
    _run(p, "Que tal como surge de la copia de poder general que acompaño, que declaro bajo juramento "
            "es fiel a su original que se encuentra vigente, soy mandataria de LIBRA COMPANIA ARGENTINA "
            "DE SEGUROS S.A., con domicilio real en Olazábal 1515, piso 5 of. 505, de la Ciudad de Buenos Aires.", bold=True)
    
    tipo_interv = d.get("tipo_intervencion", "citación en garantía")
    body(doc,
         f"En el carácter invocado, vengo a contestar la {tipo_interv} que se le ha cursado "
         "a mi conferente, solicitando su total rechazo con costas.")


def sec_objeto(doc, d, n):
    section_title(doc, "OBJETO", numero=to_roman(n))
    fecha_notif = d.get("fecha_notificacion", "[FECHA NOTIFICACIÓN — A COMPLETAR]")
    fecha_vto   = d.get("fecha_vencimiento",  "[FECHA VENCIMIENTO — A COMPLETAR]")
    tipo_interv = d.get("tipo_intervencion", "citación en garantía")
    
    comment = None
    if "[A COMPLETAR]" in fecha_notif:
        comment = "⚠️ Completar con la fecha real de la cédula de notificación. El vencimiento se calcula a 15 días hábiles (art. 346 CPCyCN para citaciones en garantía)."
    
    p = _para(doc, space_before=2, space_after=6)
    _run(p, "Que, en el carácter invocado, vengo por el presente a contestar, en legal tiempo y forma, "
            f"la {tipo_interv} cursada en estas actuaciones, notificada por cédula el ")
    _run(p, fecha_notif, underline=("[A COMPLETAR]" not in fecha_notif))
    _run(p, ", cuyo vencimiento opera el ")
    _run(p, fecha_vto, underline=("[A COMPLETAR]" not in fecha_vto))
    _run(p, ". Contestamos, asimismo, los términos de la demanda, cuyo rechazo pedimos expresamente, con costas.")
    if comment:
        add_comment(p, comment)


def sec_poliza(doc, d, n):
    poliza = d.get("poliza", {})
    numero       = poliza.get("numero",        "[N° PÓLIZA — A COMPLETAR]")
    vehiculo     = poliza.get("vehiculo",       "[VEHÍCULO — A COMPLETAR]")
    dominio      = poliza.get("dominio",        "[DOMINIO — A COMPLETAR]")
    asegurado    = poliza.get("asegurado",      "[ASEGURADO — A COMPLETAR]")
    vig_desde    = poliza.get("vigencia_desde", "[DESDE — A COMPLETAR]")
    vig_hasta    = poliza.get("vigencia_hasta", "[HASTA — A COMPLETAR]")
    limite       = poliza.get("limite",         "[LÍMITE — A COMPLETAR]")

    has_placeholders = any("[A COMPLETAR]" in str(v) for v in [numero, vehiculo, dominio, asegurado, vig_desde, vig_hasta, limite])
    comment = "⚠️ Verificar datos de póliza en sistemas internos de Libra (SISE) antes de presentar." if has_placeholders else None

    section_title(doc, "PÓLIZA", numero=to_roman(n))

    p = _para(doc, space_before=2, space_after=6)
    _run(p, f"Mi mandante emitió la póliza {numero} que amparaba -entre otros riesgos- por el de "
            f"responsabilidad civil hacia terceros, al vehículo {vehiculo} dominio {dominio}, "
            f"a nombre de {asegurado}, con una vigencia desde el {vig_desde} al {vig_hasta} "
            f"y con el límite de responsabilidad por acontecimiento establecido en las condiciones "
            f"particulares ({limite}). Adjunto a la presente copia de la póliza.")
    if comment:
        add_comment(p, comment)

    body(doc,
         "Cabe destacar que el actor carece de acción directa y autónoma contra mi representada "
         "por lo cual la responsabilidad asegurativa de mi mandante únicamente podrá hacerse efectiva "
         "en caso de condena respecto de su asegurado y en la medida del seguro (art. 118 LS).")


def sec_asume_cobertura(doc, d, n):
    poliza = d.get("poliza", {})
    vehiculo = poliza.get("vehiculo",  "[VEHÍCULO — A COMPLETAR]")
    dominio  = poliza.get("dominio",   "[DOMINIO — A COMPLETAR]")
    numero   = poliza.get("numero",    "[N° PÓLIZA — A COMPLETAR]")
    limite   = poliza.get("limite",    "[LÍMITE — A COMPLETAR]")

    section_title(doc, "ASUME COBERTURA – DENUNCIA LÍMITE – FORMULA RESERVA", numero=to_roman(n))

    body(doc,
        f"A la fecha de ocurrencia del hecho motivo de esta litis el vehículo {vehiculo} dominio {dominio}, "
        f"se hallaba asegurado por mi representada, bajo la póliza {numero}, por riesgo de "
        f"\"Responsabilidad Civil con límite\" (hasta {limite}), ello conforme la cláusula: "
        f"\"CG-RC 01.1 - Responsabilidad Civil - Riesgo Cubierto: El Asegurador se obliga a mantener "
        f"indemne al Asegurado y/o a la persona que con su autorización conduzca el vehículo objeto del "
        f"seguro (en adelante el Conductor), por cuanto deban a un tercero como consecuencia de daños "
        f"causados por ese vehículo o por la carga que transporte en condiciones reglamentarias, por "
        f"hechos acaecidos en el plazo convenido, en razón de la responsabilidad civil que pueda resultar "
        f"a cargo de ellos. El Asegurador asume esta obligación únicamente en favor del Asegurado y del "
        f"Conductor, hasta la suma máxima por acontecimiento, establecida en el Frente de Póliza por daños "
        f"corporales a personas, sean estas transportadas o no transportadas y por daños materiales, hasta "
        f"el monto máximo allí establecido para cada acontecimiento sin que los mismos puedan ser excedidos "
        f"por el conjunto de indemnizaciones que provengan de un mismo hecho generador...\" "
        f"En consecuencia, la eventual responsabilidad de mi mandante se encuentra estricta y taxativamente "
        f"limitada a los alcances, condiciones y montos máximos pactados contractualmente, los cuales "
        f"resultan plenamente oponibles tanto al asegurado como a los terceros reclamantes.")

    body(doc,
        "Manifiesto que mi mandante subordina la asunción de la responsabilidad derivada de la póliza "
        "contratada, así como de su posterior actuación en este juicio, a la efectiva traba de la litis "
        "con respecto al asegurado, toda vez que el actor carece de acción en \"forma autónoma\" con "
        "relación a mi mandante para promover esta demanda.")

    body(doc,
        "Asimismo, mi representada formula expresa reserva de invocar alguna de las causas de "
        "\"exclusión de cobertura\" previstas en la póliza y la Ley 17.418 con sus consecuencias jurídicas, "
        "si con posterioridad a éste responde llegaran a su conocimiento hechos, circunstancias y/o "
        "elementos no denunciados por su asegurado que hubieran obstado a esta presentación (Cláusula 3), "
        "haciendo expresa reserva de los derechos que pudieran derivarse de dicha circunstancia.")


def sec_negativa_general(doc, n):
    section_title(doc, "NEGATIVA GENERAL", numero=to_roman(n))
    body(doc,
        "Conforme lo dispone el art. 356 CPCCN, niego todos y cada uno de los hechos y el derecho "
        "invocado por la actora, como así también la autenticidad de toda la documentación que no sea "
        "expresamente reconocida en este responde.")


def sec_negativas_especificas(doc, d, n):
    negativas = d.get("negativas_especificas", [])
    section_title(doc, "NEGATIVAS ESPECÍFICAS", numero=to_roman(n))

    if not negativas:
        p = _para(doc, left_indent=0.5)
        _run(p, "[NEGATIVAS ESPECÍFICAS — A COMPLETAR con los hechos de la demanda]")
        add_comment(p, "⚠️ Completar con las negativas específicas de los hechos de la demanda.")
        return

    intro = _para(doc, space_before=2, space_after=4)
    _run(intro, "En particular, niego:")

    for item in negativas:
        if isinstance(item, dict):
            num_h = item.get("hecho_numero", "")
            hecho = item.get("hecho_original", item.get("texto", ""))
        else:
            num_h, hecho = "", str(item)

        prefix = f"{num_h}. " if num_h else ""
        p = _para(doc, space_after=4)
        _run(p, f"{prefix}Que {hecho}")


def sec_excepciones(doc, d, n):
    excepciones = d.get("excepciones_previas", [])
    if not excepciones:
        return n
    section_title(doc, "EXCEPCIONES PREVIAS", numero=to_roman(n))
    for exc in excepciones:
        if isinstance(exc, dict):
            tipo = exc.get("tipo", "")
            fund = exc.get("fundamento", "")
            has_placeholder = "[A COMPLETAR]" in fund
            p = _para(doc, left_indent=0.5, space_after=6)
            _run(p, f"{tipo}: {fund}")
            if has_placeholder:
                add_comment(p, f"⚠️ {fund}")
        else:
            body(doc, str(exc), indent=0.5)
    return n + 1


def sec_defensas_fondo(doc, d, n):
    defensas = d.get("defensas_fondo", [])
    if not defensas:
        return n
    section_title(doc, "DEFENSA DE FONDO", numero=to_roman(n))
    for df in defensas:
        if isinstance(df, dict):
            nombre = df.get("nombre", "")
            texto  = df.get("texto", df.get("fundamento", ""))
            p = _para(doc, space_before=4, space_after=4)
            if nombre:
                _run(p, f"{nombre}. ", bold=True)
            _run(p, str(texto))
        else:
            body(doc, str(df))
    return n + 1


def sec_prueba(doc, d, n):
    section_title(doc, "PRUEBA", numero=to_roman(n))
    prueba = d.get("ofrecimiento_prueba", {})
    tipos  = [
        ("documental",  "A) PRUEBA DOCUMENTAL"),
        ("informativa", "B) PRUEBA INFORMATIVA"),
        ("pericial",    "C) PRUEBA PERICIAL"),
        ("testimonial", "D) PRUEBA TESTIMONIAL"),
        ("confesional", "E) PRUEBA CONFESIONAL"),
    ]
    for key, label in tipos:
        items = prueba.get(key, [])
        if not items:
            continue
        p = _para(doc, space_before=6, space_after=3)
        _run(p, label, bold=True, underline=True)
        if isinstance(items, list):
            for item in items:
                desc = item.get("descripcion", str(item)) if isinstance(item, dict) else str(item)
                body(doc, desc, indent=0.5, space_after=3)
        elif isinstance(items, str):
            body(doc, items, indent=0.5)


def sec_oponibilidad(doc, d, n):
    opon = d.get("oponibilidad", "")
    if not opon:
        return n
    section_title(doc, "OPONIBILIDAD DE LAS CLÁUSULAS DE SEGURO", numero=to_roman(n))
    # Split into paragraphs if it's a long text
    if isinstance(opon, str):
        for para_text in opon.split("\n\n"):
            if para_text.strip():
                body(doc, para_text.strip())
    elif isinstance(opon, list):
        for item in opon:
            body(doc, str(item))
    return n + 1


def sec_defensa_juicio(doc, d, n):
    dj = d.get("defensa_juicio_asegurado", "")
    if not dj:
        return n
    section_title(doc, "DEFENSA EN JUICIO DEL ASEGURADO", numero=to_roman(n))
    body(doc, str(dj))
    return n + 1


def sec_impugnacion_documental(doc, d, n):
    docs_list = d.get("impugnacion_documental", [])
    if not docs_list:
        return n
    section_title(doc, "DESCONOCE DOCUMENTAL", numero=to_roman(n))
    body(doc,
        "En los términos del artículo 356 inciso 1° del Código Procesal Civil y Comercial, "
        "mi parte desconoce expresa y formalmente, por no constarle su autenticidad, veracidad, "
        "origen, integridad, fecha cierta ni correspondencia con el hecho de autos, toda la "
        "documentación acompañada por la actora. En particular, niego y desconozco:")
    for item in docs_list:
        if isinstance(item, dict):
            desc = item.get("documento", item.get("descripcion", str(item)))
        else:
            desc = str(item)
        body(doc, f"• {desc}", indent=0.5, space_after=3)
    return n + 1


def sec_impugnacion_montos(doc, d, n):
    rubros = d.get("impugnacion_montos", [])
    if not rubros:
        return n
    section_title(doc, "IMPUGNACIÓN DE MONTOS RECLAMADOS", numero=to_roman(n))
    body(doc,
        "Sin perjuicio de la negativa general, y para el hipotético e improbable supuesto de "
        "que se admitiera parcialmente la pretensión, esta parte impugna los montos reclamados "
        "por los siguientes rubros:")
    for rubro in rubros:
        if isinstance(rubro, dict):
            nombre = rubro.get("rubro", "")
            monto = rubro.get("monto_reclamado", "")
            imp = rubro.get("impugnacion", "")
            p = _para(doc, space_before=4, space_after=4, left_indent=0.5)
            _run(p, f"{nombre}", bold=True)
            if monto:
                _run(p, f" ({monto})")
            _run(p, f": {imp}")
        else:
            body(doc, str(rubro), indent=0.5)
    return n + 1


def sec_reserva_federal(doc, n):
    section_title(doc, "RESERVA DEL CASO FEDERAL", numero=to_roman(n))
    body(doc,
        "Para el hipotético caso de que V.S. no haga lugar a las defensas articuladas, se violarían "
        "derechos y garantías constitucionales, por lo que esta parte formula expresa reserva para "
        "ocurrir ante la Excma. Corte Suprema de Justicia de la Nación conforme el art. 14 de la ley 48.")


def sec_petitorio(doc, d, n):
    section_title(doc, "PETITORIO", numero=to_roman(n))
    pet = d.get("petitorio", {})
    texto = pet.get("texto", "") if isinstance(pet, dict) else str(pet)
    if texto:
        body(doc, texto)
    else:
        body(doc,
            "Por lo expuesto, solicito: a) Se me tenga por presentado, por parte en el carácter invocado "
            "y por constituido el domicilio procesal indicado; b) Se rechace en todas sus partes la "
            "demanda/citación en garantía con expresa imposición de costas; c) Se provea la prueba "
            "ofrecida en la oportunidad procesal correspondiente.")

    p = _para(doc, space_before=16, space_after=4)
    _run(p, "Proveer así,")
    p2 = _para(doc, space_after=4)
    _run(p2, "Será Justicia.", bold=True)


# ── Main builder ──────────────────────────────────────────────────────────────

def build_contestacion(data: dict, output_path: str):
    doc = Document()
    setup_document(doc)

    sec_encabezado(doc, data)

    n = 1
    sec_personeria(doc, data, n); n += 1
    sec_objeto(doc, data, n); n += 1
    sec_poliza(doc, data, n); n += 1
    sec_asume_cobertura(doc, data, n); n += 1
    # Oponibilidad (si existe)
    if data.get("oponibilidad"):
        n = sec_oponibilidad(doc, data, n)

    # Defensa en juicio del asegurado (si existe)
    if data.get("defensa_juicio_asegurado"):
        n = sec_defensa_juicio(doc, data, n)

    sec_negativa_general(doc, n); n += 1

    # Impugnación documental (antes de negativas específicas)
    if data.get("impugnacion_documental"):
        n = sec_impugnacion_documental(doc, data, n)

    sec_negativas_especificas(doc, data, n); n += 1

    if data.get("excepciones_previas"):
        n = sec_excepciones(doc, data, n)

    if data.get("defensas_fondo"):
        n = sec_defensas_fondo(doc, data, n)

    # Impugnación de montos (después de defensas)
    if data.get("impugnacion_montos"):
        n = sec_impugnacion_montos(doc, data, n)

    sec_prueba(doc, data, n); n += 1
    sec_reserva_federal(doc, n); n += 1
    sec_petitorio(doc, data, n)

    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Uso: python3 build_contestacion.py <datos.json> <output.docx>")
        sys.exit(1)
    with open(sys.argv[1], encoding="utf-8") as f:
        data = json.load(f)
    out = build_contestacion(data, sys.argv[2])
    print(f"DOCX generado: {out}")
