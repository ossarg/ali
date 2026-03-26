#!/usr/bin/env python3
"""Convert pipeline v2 MD borrador to DOCX — applies style guide formatting."""
import sys
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_para_indent(paragraph, first_line_cm=1.0):
    pf = paragraph.paragraph_format
    pf.first_line_indent = Cm(first_line_cm)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Extract JESS borrador section (after the header block, before LOU)
    jess_match = re.search(r'## JESS — Borrador de Contestación\n\n(.*?)(?=\n---\n\n## LOU)', content, re.DOTALL)
    if not jess_match:
        jess_match = re.search(r'## JESS.*?\n\n(CONTESTA DEMANDA.*?)(?=\n---\n\n## LOU|\Z)', content, re.DOTALL)
    
    if jess_match:
        borrador_text = jess_match.group(1).strip()
        # Remove the italic note at start if present
        borrador_text = re.sub(r'^\*\(Aplicando.*?\)\*\n\n', '', borrador_text)
        borrador_text = re.sub(r'^\*\(.*?\)\*\n\n', '', borrador_text)
    else:
        borrador_text = content

    # Also get Lou review report
    lou_match = re.search(r'## LOU — Quality Review\n\n(.*?)(?=\n---\n\n## PIPELINE|\Z)', content, re.DOTALL)
    lou_text = lou_match.group(1).strip() if lou_match else ""

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Heading styles
    for i in [1, 2]:
        h = doc.styles[f'Heading {i}']
        h.font.name = 'Times New Roman'
        h.font.size = Pt(12)
        h.font.bold = True
        h.font.color.rgb = RGBColor(0, 0, 0)

    def add_body_para(text, bold_parts=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.0)
        p.paragraph_format.space_after = Pt(0)

        # Split on COMPLETAR/NOTA markers and bold markers
        parts = re.split(r'(\[COMPLETAR[^\]]*\]|\[NOTA INTERNA[^\]]*\]|\*\*[^*]+\*\*)', text)
        for part in parts:
            if part.startswith('[COMPLETAR'):
                run = p.add_run(part)
                run.bold = True
                run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            elif part.startswith('[NOTA INTERNA'):
                run = p.add_run(part)
                run.italic = True
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            elif part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part:
                p.add_run(part)
        return p

    lines = borrador_text.split('\n')
    i = 0
    last_was_blank = False

    while i < len(lines):
        line = lines[i].rstrip()

        # Skip markdown headers
        if line.startswith('## ') or line.startswith('# '):
            i += 1
            continue

        # Encabezado del borrador
        if re.match(r'^(CONTESTA DEMANDA|Expediente:|Caratulado:|Tribunal:|Generado por:|Fecha:|Estado:)', line):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(11)
            i += 1
            continue

        # Section headers ALL CAPS (PERSONERÍA, OBJETO, etc.)
        stripped = line.strip()
        if stripped and stripped == stripped.upper() and len(stripped) > 5 and not stripped.startswith('['):
            # Check if it's a real section heading (not a sentence fragment)
            if re.match(r'^[A-ZÁÉÍÓÚÑ\s\-\–]+$', stripped) and len(stripped.split()) <= 10:
                p = doc.add_heading(stripped, level=1)
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(4)
                i += 1
                continue

        # Sub-headings (lines like "Reparación del rodado — $12.187.600")
        if re.match(r'^[A-ZÁÉÍÓÚÑ][a-záéíóúñ].*—.*\$', line) or re.match(r'^RUBROS DE ', line):
            p = doc.add_heading(line, level=2)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Closing line
        if '**SERÁ JUSTICIA.**' in line or line.strip() == 'SERÁ JUSTICIA.':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('Proveer de conformidad,')
            run.bold = True
            doc.add_paragraph()
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = p2.add_run('SERÁ JUSTICIA.')
            run2.bold = True
            i += 1
            continue

        # Empty line
        if not line.strip():
            last_was_blank = True
            i += 1
            continue

        # Regular body paragraph
        add_body_para(line)
        last_was_blank = False
        i += 1

    # Add Lou review as appendix
    if lou_text:
        doc.add_page_break()
        doc.add_heading('REVISIÓN DE CALIDAD — LOU', level=1)
        for line in lou_text.split('\n'):
            line = line.rstrip()
            if not line:
                continue
            if line.startswith('**') and line.endswith('**'):
                p = doc.add_paragraph()
                run = p.add_run(line[2:-2])
                run.bold = True
            else:
                # Parse inline bold
                parts = re.split(r'(\*\*[^*]+\*\*)', line)
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        p.add_run(part[2:-2]).bold = True
                    elif part.startswith('[NOTA INTERNA'):
                        run = p.add_run(part)
                        run.italic = True
                        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                    elif part:
                        p.add_run(part)

    doc.save(docx_path)
    print(f"Saved: {docx_path}")

if __name__ == '__main__':
    md_to_docx(sys.argv[1], sys.argv[2])
